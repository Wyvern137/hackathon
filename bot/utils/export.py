"""
Утилиты для экспорта контента
"""
import logging
import csv
import json
from datetime import datetime, date, timedelta
from pathlib import Path
from typing import List, Optional, Dict, Any
from bot.database.models import ContentHistory, ContentPlan
from bot.database.database import get_db

try:
    from docx import Document
    docx_available = True
except ImportError:
    docx_available = False
    logging.warning("python-docx не установлен, экспорт в DOCX будет недоступен")

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import inch
    reportlab_available = True
except ImportError:
    reportlab_available = False
    logging.warning("reportlab не установлен, экспорт в PDF будет недоступен")

try:
    import openpyxl
    from openpyxl import Workbook
    openpyxl_available = True
except ImportError:
    openpyxl_available = False
    logging.warning("openpyxl не установлен, экспорт в Excel будет недоступен")

try:
    from icalendar import Calendar, Event
    icalendar_available = True
except ImportError:
    icalendar_available = False
    logging.warning("icalendar не установлен, экспорт в iCal будет недоступен")

logger = logging.getLogger(__name__)

EXPORT_DIR = Path("data/exports")
EXPORT_DIR.mkdir(parents=True, exist_ok=True)


async def export_history_to_txt(user_id: int, limit: Optional[int] = None) -> Optional[Path]:
    """
    Экспортирует историю контента в текстовый файл
    
    Args:
        user_id: ID пользователя
        limit: Максимальное количество записей (None = все)
    
    Returns:
        Path к файлу или None
    """
    try:
        with get_db() as db:
            query = db.query(ContentHistory).filter(
                ContentHistory.user_id == user_id
            ).order_by(ContentHistory.generated_at.desc())
            
            if limit:
                query = query.limit(limit)
            
            history_items = query.all()
        
        if not history_items:
            return None
        
        # Создаем файл
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"history_{user_id}_{timestamp}.txt"
        file_path = EXPORT_DIR / filename
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write("=" * 80 + "\n")
            f.write("ИСТОРИЯ КОНТЕНТА\n")
            f.write("=" * 80 + "\n\n")
            
            for i, item in enumerate(history_items, 1):
                date_str = item.generated_at.strftime("%d.%m.%Y %H:%M")
                f.write(f"Запись {i}\n")
                f.write("-" * 80 + "\n")
                f.write(f"Дата: {date_str}\n")
                f.write(f"Тип: {item.content_type}\n")
                
                if item.content_type == "text":
                    content_data = item.content_data if isinstance(item.content_data, dict) else {}
                    text = content_data.get("text", str(item.content_data))
                    hashtags = content_data.get("hashtags", [])
                    
                    f.write(f"\nТекст:\n{text}\n")
                    
                    if hashtags:
                        f.write(f"\nХештеги: {' '.join(hashtags)}\n")
                
                f.write("\n" + "=" * 80 + "\n\n")
        
        logger.info(f"Экспортировано {len(history_items)} записей в {file_path}")
        return file_path
    
    except Exception as e:
        logger.exception(f"Ошибка при экспорте истории: {e}")
        return None


async def export_content_plan_to_csv(user_id: int, plan_id: Optional[int] = None) -> Optional[Path]:
    """
    Экспортирует контент-план в CSV файл
    
    Args:
        user_id: ID пользователя
        plan_id: ID плана (None = все активные планы)
    
    Returns:
        Path к файлу или None
    """
    try:
        with get_db() as db:
            query = db.query(ContentPlan).filter(
                ContentPlan.user_id == user_id
            )
            
            if plan_id:
                query = query.filter(ContentPlan.id == plan_id)
            else:
                query = query.filter(ContentPlan.is_active == True)
            
            plans = query.all()
        
        if not plans:
            return None
        
        # Создаем CSV файл
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"content_plans_{user_id}_{timestamp}.csv"
        file_path = EXPORT_DIR / filename
        
        with open(file_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            
            # Заголовки
            writer.writerow([
                "План", "Начало", "Окончание", "Частота", "Дни недели", "Время", "Темы"
            ])
            
            for plan in plans:
                schedule = plan.schedule if isinstance(plan.schedule, dict) else {}
                
                days = schedule.get("days", [])
                days_names = ['пн', 'вт', 'ср', 'чт', 'пт', 'сб', 'вс']
                days_str = ', '.join([days_names[d-1] for d in days]) if days else "не указано"
                
                writer.writerow([
                    plan.plan_name,
                    plan.start_date.strftime("%d.%m.%Y"),
                    plan.end_date.strftime("%d.%m.%Y"),
                    f"{plan.frequency} раз в неделю",
                    days_str,
                    schedule.get("time", "не указано"),
                    schedule.get("topics", "не указано")
                ])
        
        logger.info(f"Экспортировано {len(plans)} планов в {file_path}")
        return file_path
    
    except Exception as e:
        logger.exception(f"Ошибка при экспорте контент-плана: {e}")
        return None


async def export_texts_to_csv(user_id: int, limit: Optional[int] = None) -> Optional[Path]:
    """
    Экспортирует тексты в CSV файл
    
    Args:
        user_id: ID пользователя
        limit: Максимальное количество записей
    
    Returns:
        Path к файлу или None
    """
    try:
        with get_db() as db:
            query = db.query(ContentHistory).filter(
                ContentHistory.user_id == user_id,
                ContentHistory.content_type == "text"
            ).order_by(ContentHistory.generated_at.desc())
            
            if limit:
                query = query.limit(limit)
            
            texts = query.all()
        
        if not texts:
            return None
        
        # Создаем CSV файл
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"texts_{user_id}_{timestamp}.csv"
        file_path = EXPORT_DIR / filename
        
        with open(file_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            
            # Заголовки
            writer.writerow(["Дата", "Текст", "Хештеги", "Стиль", "Тип"])
            
            for text_item in texts:
                content_data = text_item.content_data if isinstance(text_item.content_data, dict) else {}
                text = content_data.get("text", str(text_item.content_data))
                hashtags = ' '.join(content_data.get("hashtags", []))
                style = content_data.get("style", "не указан")
                text_type = content_data.get("type", "обычный")
                
                writer.writerow([
                    text_item.generated_at.strftime("%d.%m.%Y %H:%M"),
                    text,
                    hashtags,
                    style,
                    text_type
                ])
        
        logger.info(f"Экспортировано {len(texts)} текстов в {file_path}")
        return file_path
    
    except Exception as e:
        logger.exception(f"Ошибка при экспорте текстов: {e}")
        return None


async def export_to_docx(user_id: int, limit: Optional[int] = None) -> Optional[Path]:
    """
    Экспортирует тексты в DOCX документ
    
    Args:
        user_id: ID пользователя
        limit: Максимальное количество записей
    
    Returns:
        Path к файлу или None
    """
    if not docx_available:
        logger.warning("python-docx не установлен, экспорт в DOCX недоступен")
        return None
    
    try:
        with get_db() as db:
            query = db.query(ContentHistory).filter(
                ContentHistory.user_id == user_id,
                ContentHistory.content_type == "text"
            ).order_by(ContentHistory.generated_at.desc())
            
            if limit:
                query = query.limit(limit)
            
            texts = query.all()
        
        if not texts:
            return None
        
        # Создаем документ
        doc = Document()
        doc.add_heading('История контента', 0)
        
        for i, text_item in enumerate(texts, 1):
            date_str = text_item.generated_at.strftime("%d.%m.%Y %H:%M")
            doc.add_heading(f'Запись {i} - {date_str}', level=1)
            
            content_data = text_item.content_data if isinstance(text_item.content_data, dict) else {}
            text = content_data.get("text", str(text_item.content_data))
            hashtags = content_data.get("hashtags", [])
            
            doc.add_paragraph(text)
            
            if hashtags:
                doc.add_paragraph(f"Хештеги: {' '.join(hashtags)}")
        
        # Сохраняем файл
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"history_{user_id}_{timestamp}.docx"
        file_path = EXPORT_DIR / filename
        doc.save(file_path)
        
        logger.info(f"Экспортировано {len(texts)} текстов в DOCX: {file_path}")
        return file_path
    
    except Exception as e:
        logger.exception(f"Ошибка при экспорте в DOCX: {e}")
        return None


async def export_to_pdf(user_id: int, limit: Optional[int] = None) -> Optional[Path]:
    """
    Экспортирует тексты в PDF документ
    
    Args:
        user_id: ID пользователя
        limit: Максимальное количество записей
    
    Returns:
        Path к файлу или None
    """
    if not reportlab_available:
        logger.warning("reportlab не установлен, экспорт в PDF недоступен")
        return None
    
    try:
        with get_db() as db:
            query = db.query(ContentHistory).filter(
                ContentHistory.user_id == user_id,
                ContentHistory.content_type == "text"
            ).order_by(ContentHistory.generated_at.desc())
            
            if limit:
                query = query.limit(limit)
            
            texts = query.all()
        
        if not texts:
            return None
        
        # Создаем PDF
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"history_{user_id}_{timestamp}.pdf"
        file_path = EXPORT_DIR / filename
        
        c = canvas.Canvas(str(file_path), pagesize=A4)
        width, height = A4
        
        y_position = height - 50
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, y_position, "История контента")
        y_position -= 40
        
        c.setFont("Helvetica", 12)
        
        for i, text_item in enumerate(texts, 1):
            if y_position < 100:
                c.showPage()
                y_position = height - 50
            
            date_str = text_item.generated_at.strftime("%d.%m.%Y %H:%M")
            c.setFont("Helvetica-Bold", 14)
            c.drawString(50, y_position, f"Запись {i} - {date_str}")
            y_position -= 25
            
            content_data = text_item.content_data if isinstance(text_item.content_data, dict) else {}
            text = content_data.get("text", str(content_data))
            hashtags = content_data.get("hashtags", [])
            
            c.setFont("Helvetica", 11)
            # Простой перенос текста
            lines = text.split('\n')
            for line in lines:
                if y_position < 100:
                    c.showPage()
                    y_position = height - 50
                
                # Разбиваем длинные строки
                if len(line) > 80:
                    words = line.split()
                    current_line = ""
                    for word in words:
                        if len(current_line + word) > 80:
                            if current_line:
                                c.drawString(50, y_position, current_line)
                                y_position -= 15
                            current_line = word + " "
                        else:
                            current_line += word + " "
                    if current_line:
                        c.drawString(50, y_position, current_line)
                        y_position -= 15
                else:
                    c.drawString(50, y_position, line)
                    y_position -= 15
            
            if hashtags:
                c.setFont("Helvetica-Oblique", 10)
                hashtags_text = f"Хештеги: {' '.join(hashtags)}"
                c.drawString(50, y_position, hashtags_text)
                y_position -= 20
            
            y_position -= 10
        
        c.save()
        
        logger.info(f"Экспортировано {len(texts)} текстов в PDF: {file_path}")
        return file_path
    
    except Exception as e:
        logger.exception(f"Ошибка при экспорте в PDF: {e}")
        return None


async def export_plan_to_excel(user_id: int, plan_id: Optional[int] = None) -> Optional[Path]:
    """
    Экспортирует контент-план в Excel файл
    
    Args:
        user_id: ID пользователя
        plan_id: ID плана (None = все активные планы)
    
    Returns:
        Path к файлу или None
    """
    if not openpyxl_available:
        logger.warning("openpyxl не установлен, экспорт в Excel недоступен")
        return None
    
    try:
        with get_db() as db:
            query = db.query(ContentPlan).filter(
                ContentPlan.user_id == user_id
            )
            
            if plan_id:
                query = query.filter(ContentPlan.id == plan_id)
            else:
                query = query.filter(ContentPlan.is_active == True)
            
            plans = query.all()
        
        if not plans:
            return None
        
        # Создаем Excel файл
        wb = Workbook()
        ws = wb.active
        ws.title = "Контент-планы"
        
        # Заголовки
        headers = ["План", "Начало", "Окончание", "Частота", "Дни недели", "Время", "Темы", "Статус"]
        ws.append(headers)
        
        # Данные
        for plan in plans:
            schedule = plan.schedule if isinstance(plan.schedule, dict) else {}
            days = schedule.get("days", [])
            days_names = ['пн', 'вт', 'ср', 'чт', 'пт', 'сб', 'вс']
            days_str = ', '.join([days_names[d-1] for d in days]) if days else "не указано"
            
            ws.append([
                plan.plan_name,
                plan.start_date.strftime("%d.%m.%Y"),
                plan.end_date.strftime("%d.%m.%Y"),
                f"{plan.frequency} раз в неделю",
                days_str,
                schedule.get("time", "не указано"),
                schedule.get("topics", "не указано"),
                "Активен" if plan.is_active else "Неактивен"
            ])
        
        # Сохраняем файл
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"content_plans_{user_id}_{timestamp}.xlsx"
        file_path = EXPORT_DIR / filename
        wb.save(file_path)
        
        logger.info(f"Экспортировано {len(plans)} планов в Excel: {file_path}")
        return file_path
    
    except Exception as e:
        logger.exception(f"Ошибка при экспорте в Excel: {e}")
        return None


async def export_to_ical(user_id: int, plan_id: Optional[int] = None) -> Optional[Path]:
    """
    Экспортирует контент-план в iCal формат (.ics)
    
    Args:
        user_id: ID пользователя
        plan_id: ID плана (None = все активные планы)
    
    Returns:
        Path к файлу или None
    """
    if not icalendar_available:
        logger.warning("icalendar не установлен, экспорт в iCal недоступен")
        return None
    
    try:
        with get_db() as db:
            query = db.query(ContentPlan).filter(
                ContentPlan.user_id == user_id,
                ContentPlan.is_active == True
            )
            
            if plan_id:
                query = query.filter(ContentPlan.id == plan_id)
            
            plans = query.all()
        
        if not plans:
            return None
        
        # Создаем календарь
        cal = Calendar()
        cal.add('prodid', '-//NKO Bot Content Plan//EN')
        cal.add('version', '2.0')
        
        # Добавляем события для каждого плана
        for plan in plans:
            schedule = plan.schedule if isinstance(plan.schedule, dict) else {}
            days = schedule.get("days", [])
            topics = schedule.get("topics", [])
            time_str = schedule.get("time", "12:00")
            
            # Парсим время
            try:
                hour, minute = map(int, time_str.split(':'))
            except:
                hour, minute = 12, 0
            
            # Генерируем события для всех дней публикации
            current_date = plan.start_date
            event_num = 0
            
            while current_date <= plan.end_date:
                weekday = current_date.weekday() + 1  # 1=понедельник, 7=воскресенье
                
                if weekday in days:
                    event = Event()
                    event.add('summary', f"{plan.plan_name} - Публикация поста")
                    
                    # Тема поста
                    topic = topics[event_num % len(topics)] if topics else "Пост для НКО"
                    event.add('description', f"Тема поста: {topic}")
                    
                    # Дата и время
                    event.add('dtstart', datetime.combine(current_date, datetime.min.time().replace(hour=hour, minute=minute)))
                    event.add('dtend', datetime.combine(current_date, datetime.min.time().replace(hour=hour+1, minute=minute)))
                    event.add('dtstamp', datetime.now())
                    
                    cal.add_component(event)
                    event_num += 1
                
                current_date += timedelta(days=1)
        
        # Сохраняем файл
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"content_plan_{user_id}_{timestamp}.ics"
        file_path = EXPORT_DIR / filename
        
        with open(file_path, 'wb') as f:
            f.write(cal.to_ical())
        
        logger.info(f"Экспортирован контент-план в iCal: {file_path}")
        return file_path
    
    except Exception as e:
        logger.exception(f"Ошибка при экспорте в iCal: {e}")
        return None


async def create_images_archive(user_id: int, limit: Optional[int] = None) -> Optional[Path]:
    """
    Создает ZIP-архив с изображениями пользователя
    
    Args:
        user_id: ID пользователя
        limit: Максимальное количество изображений
    
    Returns:
        Path к архиву или None
    """
    try:
        import zipfile
        from bot.config import config
        
        with get_db() as db:
            query = db.query(ContentHistory).filter(
                ContentHistory.user_id == user_id,
                ContentHistory.content_type == "image"
            ).order_by(ContentHistory.generated_at.desc())
            
            if limit:
                query = query.limit(limit)
            
            images = query.all()
        
        if not images:
            return None
        
        # Создаем ZIP архив
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"images_{user_id}_{timestamp}.zip"
        archive_path = EXPORT_DIR / filename
        
        with zipfile.ZipFile(archive_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for i, image_item in enumerate(images, 1):
                content_data = image_item.content_data if isinstance(image_item.content_data, dict) else {}
                file_path = content_data.get("file_path") or content_data.get("path")
                
                if file_path and Path(file_path).exists():
                    # Добавляем файл в архив
                    arcname = f"image_{i}_{Path(file_path).name}"
                    zipf.write(file_path, arcname)
        
        logger.info(f"Создан архив с {len(images)} изображениями: {archive_path}")
        return archive_path
    
    except Exception as e:
        logger.exception(f"Ошибка при создании архива изображений: {e}")
        return None


async def batch_export(
    user_id: int,
    content_types: List[str] = ["text", "image"],
    formats: List[str] = ["txt"],
    period_days: Optional[int] = None
) -> Dict[str, Optional[Path]]:
    """
    Пакетный экспорт контента за период
    
    Args:
        user_id: ID пользователя
        content_types: Типы контента для экспорта
        formats: Форматы для экспорта (txt, docx, pdf, csv, excel)
        period_days: Период в днях (None = весь период)
    
    Returns:
        Dict с путями к файлам
    """
    results = {}
    
    try:
        start_date = datetime.now() - timedelta(days=period_days) if period_days else None
        
        for content_type in content_types:
            if content_type == "text":
                for fmt in formats:
                    if fmt == "txt":
                        results["texts_txt"] = await export_history_to_txt(user_id)
                    elif fmt == "docx" and docx_available:
                        results["texts_docx"] = await export_to_docx(user_id)
                    elif fmt == "pdf" and reportlab_available:
                        results["texts_pdf"] = await export_to_pdf(user_id)
                    elif fmt == "csv":
                        results["texts_csv"] = await export_texts_to_csv(user_id)
                    elif fmt == "excel" and openpyxl_available:
                        # Можно использовать CSV как альтернативу Excel
                        results["texts_excel"] = await export_texts_to_csv(user_id)
            
            elif content_type == "image":
                results["images_zip"] = await create_images_archive(user_id)
        
        return results
    
    except Exception as e:
        logger.exception(f"Ошибка при пакетном экспорте: {e}")
        return results

